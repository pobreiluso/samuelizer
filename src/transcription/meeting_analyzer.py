from docx import Document
import logging
import os
import openai
from typing import Optional, List, Dict, Any
from src.transcription.exceptions import AnalysisError
from src.interfaces import TranscriptionService, TextAnalysisModelInterface
from .templates import PromptTemplates
from src.models.model_factory import ModelProviderFactory

logger = logging.getLogger(__name__)

from src.transcription.text_preprocessor import TextPreprocessor

class AnalysisClient:
    """
    Cliente de análisis que utiliza el proveedor de modelos configurado.
    Maneja diferentes tipos de modelos (chat y completions).
    """
    def __init__(self, provider: Optional[TextAnalysisModelInterface] = None, 
                 provider_name: str = "openai", api_key: Optional[str] = None,
                 model_id: str = "gpt-3.5-turbo"):
        """
        Inicializa el cliente de análisis
        
        Args:
            provider: Proveedor de modelos preconfigurado (opcional)
            provider_name: Nombre del proveedor a utilizar si no se proporciona uno
            api_key: Clave API para el proveedor (opcional)
            model_id: Identificador del modelo a utilizar (opcional)
        """
        self.provider = provider
        if not self.provider:
            self.provider = ModelProviderFactory.get_analysis_model(
                provider_name, api_key=api_key
            )
        self.provider_name = provider_name
        self.model_id = model_id
        
        # Configurar OpenAI API key si se proporciona
        if api_key and provider_name.lower() == "openai":
            import os
            os.environ["OPENAI_API_KEY"] = api_key

    def analyze(self, messages: List[Dict[str, str]], model_id: str = None, **kwargs) -> str:
        """
        Analiza un texto utilizando el proveedor configurado
        
        Args:
            messages: Lista de mensajes en formato compatible con el modelo
            model_id: Identificador del modelo a utilizar (opcional, usa self.model_id si no se proporciona)
            **kwargs: Parámetros adicionales para el modelo
            
        Returns:
            str: Resultado del análisis
        """
        # Usar el modelo_id pasado como parámetro o el almacenado en la instancia
        model_to_use = model_id or self.model_id
        
        # Limitar el tamaño de los mensajes para evitar errores
        max_content_length = 15000  # Ajustar según sea necesario
        for i, message in enumerate(messages):
            if "content" in message and len(message["content"]) > max_content_length:
                logger.warning(f"Mensaje demasiado largo ({len(message['content'])} caracteres). Truncando a {max_content_length} caracteres.")
                messages[i]["content"] = message["content"][:max_content_length]
        
        # Si estamos usando OpenAI directamente, manejar diferentes tipos de modelos
        if self.provider_name.lower() == "openai":
            return self._analyze_with_openai(messages, model_to_use, **kwargs)
        
        # Usar el proveedor configurado
        return self.provider.analyze(messages, model_to_use, **kwargs)
        
    def _analyze_with_openai(self, messages: List[Dict[str, str]], model_id: str, **kwargs) -> str:
        """
        Analiza texto usando OpenAI API directamente, manejando diferentes tipos de modelos
        
        Args:
            messages: Lista de mensajes en formato compatible con el modelo
            model_id: Identificador del modelo a utilizar
            **kwargs: Parámetros adicionales para el modelo
            
        Returns:
            str: Resultado del análisis
        """
        import openai
        
        # Determinar si es un modelo de chat o de completions
        chat_models = ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo", "gpt-4-1106-preview"]
        is_chat_model = any(chat_name in model_id for chat_name in chat_models)
        
        try:
            if is_chat_model:
                return self._analyze_with_chat_model(messages, model_id, **kwargs)
            else:
                try:
                    return self._analyze_with_completion_model(messages, model_id, **kwargs)
                except Exception as completion_error:
                    # Si falla el modelo de completions, intentar con un modelo de chat como fallback
                    logger.warning(f"Error con modelo de completions {model_id}: {completion_error}. Usando gpt-3.5-turbo como fallback.")
                    return self._analyze_with_chat_model(messages, "gpt-3.5-turbo", **kwargs)
        except Exception as e:
            logger.error(f"Error en OpenAI API: {e}")
            raise
            
    def _analyze_with_chat_model(self, messages: List[Dict[str, str]], model_id: str, **kwargs) -> str:
        """
        Analiza texto usando modelos de chat de OpenAI
        """
        import openai
        
        response = openai.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=kwargs.get('temperature', 0),
            max_tokens=kwargs.get('max_tokens', 1000)
        )
        return response.choices[0].message.content.strip()
        
    def _analyze_with_completion_model(self, messages: List[Dict[str, str]], model_id: str, **kwargs) -> str:
        """
        Analiza texto usando modelos de completions de OpenAI
        """
        import openai
        
        # Extraer el prompt de los mensajes
        prompt = ""
        for message in messages:
            role = message.get("role", "")
            content = message.get("content", "")
            if role and content:
                prompt += f"{role.upper()}: {content}\n\n"
        
        # Limitar el tamaño del prompt para evitar errores
        max_prompt_length = 4000  # Ajustar según sea necesario
        if len(prompt) > max_prompt_length:
            logger.warning(f"Prompt demasiado largo ({len(prompt)} caracteres). Truncando a {max_prompt_length} caracteres.")
            prompt = prompt[:max_prompt_length]
        
        response = openai.completions.create(
            model=model_id,
            prompt=prompt,
            temperature=kwargs.get('temperature', 0),
            max_tokens=kwargs.get('max_tokens', 1000)
        )
        return response.choices[0].text.strip()

class TemplateSelector:
    """
    Selects the appropriate template for analysis.
    """
    def __init__(self, prompt_templates, analysis_client, model_id="gpt-3.5-turbo"):
        self.prompt_templates = prompt_templates
        self.analysis_client = analysis_client
        self.model_id = model_id

    def select_template(self, text, **kwargs):
        template = self.prompt_templates.get_template("auto")
        try:
            messages = [
                {"role": "system", "content": template["system"]},
                {"role": "user", "content": template["template"].format(text=text)}
            ]
            analysis = self.analysis_client.analyze(messages, model_id=self.model_id)
            first_line = analysis.split('\n')[0].strip()
            recommended_template = first_line.split(':')[-1].strip().lower().replace('**', '').replace('*', '')
            logger.info(f"Auto-selected template: {recommended_template}")
            logger.info(f"Selection reasoning: {analysis}")
            return self.prompt_templates.get_template(recommended_template, **kwargs)
        except Exception as e:
            logger.warning(f"Error in auto-template selection: {e}. Falling back to 'summary' template.")
            return self.prompt_templates.get_template("summary", **kwargs)

class MeetingAnalyzer:
    """
    Analyzes meeting transcriptions.
    """
    def __init__(self, transcription: str, analysis_client=None, prompt_templates=None, 
                provider_name: str = "openai", api_key: str = None, model_id: str = "gpt-3.5-turbo"):
        self.text_preprocessor = TextPreprocessor()
        self.transcription = self.text_preprocessor.prepare_text(transcription)
        
        # Inicializar el cliente de análisis si no se proporciona
        if analysis_client is None:
            self.analysis_client = AnalysisClient(
                provider_name=provider_name,
                api_key=api_key,
                model_id=model_id
            )
        else:
            self.analysis_client = analysis_client
            
        self.prompt_templates = prompt_templates or PromptTemplates()
        self.template_selector = TemplateSelector(
            self.prompt_templates, 
            self.analysis_client,
            model_id=self.analysis_client.model_id
        )

    def analyze(self, template_name: str = "auto", **kwargs) -> str:
        try:
            if template_name == "auto":
                template = self.template_selector.select_template(self.transcription, **kwargs)
            else:
                template = self.prompt_templates.get_template(template_name, **kwargs)
                
            # Preparar los parámetros para el formato de la plantilla
            format_params = {'text': self.transcription}
            # Añadir los parámetros adicionales al diccionario de formato
            format_params.update(kwargs)
            
            try:
                formatted_template = template["template"].format(**format_params)
            except KeyError as e:
                logger.warning(f"Missing parameter in template: {e}. Using default values.")
                # Si falta algún parámetro, intentar con valores predeterminados
                missing_param = str(e).strip("'")
                if missing_param == 'start_date':
                    format_params['start_date'] = 'fecha no especificada'
                if missing_param == 'end_date':
                    format_params['end_date'] = 'fecha no especificada'
                if missing_param == 'channel_count':
                    format_params['channel_count'] = '0'
                formatted_template = template["template"].format(**format_params)
            
            messages = [
                {"role": "system", "content": template["system"]},
                {"role": "user", "content": formatted_template}
            ]
            return self.analysis_client.analyze(messages, **kwargs)
        except openai.AuthenticationError as e:
            logger.error(f"Error with template '{template_name}': {e}")
            raise AnalysisError(f"Authentication failed: {e}") from e
        except openai.APIError as e:
            logger.error(f"API Error: {e}")
            raise AnalysisError(f"API Error: {e}") from e
        except Exception as e:
            logger.error(f"Unexpected error during analysis: {e}")
            raise AnalysisError(f"Unexpected error: {e}") from e

    def summarize(self, **kwargs):
        return self.analyze("summary", **kwargs)

    def extract_key_points(self, **kwargs):
        return self.analyze("key_points", **kwargs)

    def extract_action_items(self, **kwargs):
        return self.analyze("action_items", **kwargs)

    def analyze_sentiment(self, **kwargs):
        return self.analyze("sentiment", **kwargs)

class DocumentManager:
    """
    Manages document creation and saving.
    """
    @staticmethod
    def create_document(content):
        doc = Document()
        for key, value in content.items():
            heading = key.replace('_', ' ').title()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        return doc

    @staticmethod
    def save_to_docx(content, filename):
        os.makedirs(os.path.dirname(os.path.abspath(filename)), exist_ok=True)
        doc = DocumentManager.create_document(content)
        doc.save(filename)
        return filename
