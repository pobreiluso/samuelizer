import os
import requests
import sys
from docx import Document
from src.utils.audio_extractor import AudioExtractor
import openai
import webbrowser
import logging
from tqdm import tqdm
from src.interfaces import TranscriptionService
from src.transcription.exceptions import (
    TranscriptionError,
    AnalysisError,
    DownloadError,
    AudioExtractionError,
    MeetingMinutesError
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('meeting_minutes.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)


class AudioFileHandler:
    """
    Handles audio file operations like getting duration and extracting segments
    """
    @staticmethod
    def get_audio_duration(audio_file_path):
        """
        Get the duration of an audio file in seconds.
        
        Args:
            audio_file_path: Path to the audio file
            
        Returns:
            float: Duration in seconds
        """
        try:
            import subprocess
            result = subprocess.run(
                ['ffprobe', '-v', 'error', '-show_entries', 'format=duration', 
                 '-of', 'default=noprint_wrappers=1:nokey=1', audio_file_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
            return float(result.stdout.strip())
        except Exception as e:
            logger.error(f"Failed to get audio duration: {e}")
            return 300.0  # Default to 5 minutes if duration can't be determined
    
    @staticmethod
    def extract_segment(audio_file_path, start_time, end_time):
        """
        Extract a segment from an audio file
        
        Args:
            audio_file_path: Path to the audio file
            start_time: Start time in seconds
            end_time: End time in seconds
            
        Returns:
            str: Path to the extracted segment
        """
        import tempfile
        import subprocess
        
        with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_file:
            temp_path = temp_file.name
        
        subprocess.run([
            'ffmpeg', '-y', '-i', audio_file_path,
            '-ss', str(start_time), '-to', str(end_time),
            '-c:a', 'copy', temp_path
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        return temp_path

class SpeakerDiarization:
    """
    Handles speaker diarization (identifying different speakers in audio)
    """
    def detect_speakers(self, audio_file_path):
        """
        Detect different speakers in an audio file
        
        Args:
            audio_file_path: Path to the audio file
            
        Returns:
            list: List of speaker segments with start and end times
        """
        try:
            logger.info("Detecting speakers in audio file...")
            # This is a placeholder implementation
            # In a real implementation, you would use a diarization library or API
            
            # Simulated speaker detection result
            segments = [
                {'speaker': 'Speaker 1', 'start': 0.0, 'end': 30.0},
                {'speaker': 'Speaker 2', 'start': 30.0, 'end': 45.0},
                {'speaker': 'Speaker 1', 'start': 45.0, 'end': 60.0}
            ]
            
            logger.info(f"Detected {len(set(s['speaker'] for s in segments))} speakers")
            return segments
        except Exception as e:
            logger.error(f"Speaker detection failed: {e}")
            # If speaker detection fails, return a single segment for the entire audio
            duration = AudioFileHandler.get_audio_duration(audio_file_path)
            return [{'speaker': 'Unknown', 'start': 0.0, 'end': duration}]

class TranscriptionClient:
    """
    Interface for transcription clients
    """
    def transcribe(self, audio_file, model, response_format="text"):
        """
        Transcribe an audio file
        
        Args:
            audio_file: Audio file object
            model: Model to use for transcription
            response_format: Format of the response
            
        Returns:
            str: Transcribed text
        """
        pass

class OpenAITranscriptionClient(TranscriptionClient):
    """
    OpenAI implementation of TranscriptionClient
    """
    def __init__(self, client=None):
        self.client = client or openai
    
    def transcribe(self, audio_file, model, response_format="text"):
        """
        Transcribe an audio file using OpenAI
        
        Args:
            audio_file: Audio file object
            model: OpenAI model to use
            response_format: Format of the response
            
        Returns:
            str: Transcribed text
        """
        return self.client.audio.transcriptions.create(
            model=model,
            file=audio_file,
            response_format=response_format
        )

class TranscriptionFileWriter:
    """
    Handles writing transcription results to files
    """
    @staticmethod
    def save_transcription(transcription, audio_file_path):
        """
        Save transcription to a text file
        
        Args:
            transcription: Transcribed text
            audio_file_path: Path to the original audio file
            
        Returns:
            str: Path to the saved transcription file
        """
        output_txt = os.path.splitext(audio_file_path)[0] + "_transcription.txt"
        with open(output_txt, 'w', encoding='utf-8') as f:
            f.write(transcription)
        logger.info(f"Transcription saved to: {output_txt}")
        return output_txt

class AudioTranscriptionService(TranscriptionService):
    """
    Service for transcribing audio files
    """
    def __init__(self, model="whisper-1", transcription_client=None, diarization_service=None):
        """
        Initialize the transcription service
        
        Args:
            model: Model to use for transcription
            transcription_client: Client for transcription API
            diarization_service: Service for speaker diarization
        """
        self.model = model
        self.transcription_client = transcription_client or OpenAITranscriptionClient()
        self.diarization_service = diarization_service or SpeakerDiarization()
        self.file_handler = AudioFileHandler()
        self.file_writer = TranscriptionFileWriter()
    
    def _transcribe_audio_segment(self, audio_file_path, start_time, end_time):
        """
        Transcribe a specific segment of an audio file
        
        Args:
            audio_file_path: Path to the audio file
            start_time: Start time in seconds
            end_time: End time in seconds
            
        Returns:
            str: Transcription of the segment
        """
        try:
            # Extract the segment
            temp_path = self.file_handler.extract_segment(audio_file_path, start_time, end_time)
            
            # Transcribe the segment
            with open(temp_path, 'rb') as audio_file:
                transcription = self.transcription_client.transcribe(
                    audio_file,
                    model=self.model
                )
            
            # Clean up
            os.unlink(temp_path)
            
            return transcription
        except Exception as e:
            logger.error(f"Error transcribing segment {start_time}-{end_time}: {e}")
            return f"[Transcription error: {str(e)}]"
    
    def transcribe(self, audio_file_path, diarization: bool = False):
        """
        Transcribe an audio file
        
        Args:
            audio_file_path: Path to the audio file
            diarization: Whether to identify different speakers
            
        Returns:
            str: Transcribed text
            
        Raises:
            TranscriptionError: If transcription fails
        """
        try:
            file_size = os.path.getsize(audio_file_path) / (1024 * 1024)  # Size in MB
            logger.info(f"Processing audio file of {file_size:.2f} MB")
            
            if diarization:
                logger.info("Diarization enabled, detecting speakers...")
                speaker_segments = self.diarization_service.detect_speakers(audio_file_path)
                full_transcript = ""
                
                for segment in speaker_segments:
                    # Transcribe the segment
                    segment_text = self._transcribe_audio_segment(
                        audio_file_path,
                        start_time=segment['start'],
                        end_time=segment['end']
                    )
                    full_transcript += f"[{segment['speaker']}]: {segment_text}\n"
                
                # Save the transcription
                self.file_writer.save_transcription(full_transcript, audio_file_path)
                return full_transcript
            else:
                with tqdm(total=100, desc="Transcribing audio", unit="%", colour='green') as pbar:
                    # Update progress: Preparation
                    pbar.update(10)
                    logger.info("Preparing file for transcription...")
                    
                    # Update progress: Opening file
                    with open(audio_file_path, 'rb') as audio_file:
                        pbar.update(20)
                        logger.info("Sending file to transcription service...")
                        
                        # Update progress: Sending to service
                        pbar.update(20)
                        
                        # Perform transcription
                        transcription = self.transcription_client.transcribe(
                            audio_file,
                            model=self.model
                        )
                        
                        # Update progress: Transcription completed
                        pbar.update(40)
                        logger.info("Transcription completed successfully")
                        
                        # Show result information
                        if transcription:
                            char_count = len(transcription)
                            word_count = len(transcription.split())
                            logger.info(f"Transcription generated: {word_count} words, {char_count} characters")
                            
                            # Save transcription to a text file
                            self.file_writer.save_transcription(transcription, audio_file_path)
                    
            return transcription
        except openai.AuthenticationError as e:
            logger.error(f"Authentication failed: {e}")
            raise TranscriptionError(f"Authentication failed: {e}") from e
        except openai.APIError as e:
            logger.error(f"API Error: {e}")
            raise TranscriptionError(f"API Error: {e}") from e
        except Exception as e:
            logger.error(f"Unexpected error during transcription: {e}")
            raise TranscriptionError(f"Unexpected error during transcription: {e}") from e


from .templates import PromptTemplates

class TextPreprocessor:
    """
    Handles text preprocessing for analysis
    """
    def __init__(self, max_chunk_size=4000):
        self.max_chunk_size = max_chunk_size
    
    def prepare_text(self, text: str) -> str:
        """
        Prepare text for analysis by cleaning and chunking if needed
        
        Args:
            text: Text to prepare
            
        Returns:
            str: Prepared text
        """
        if len(text) > self.max_chunk_size:
            # Split into semantic chunks (at paragraph or sentence boundaries)
            chunks = [text[i:i+self.max_chunk_size] for i in range(0, len(text), self.max_chunk_size)]
            return "\n\n".join(chunks)
        return text

class AnalysisClient:
    """
    Interface for analysis clients
    """
    def analyze(self, messages, model="gpt-4-1106-preview", temperature=0):
        """
        Analyze text using a language model
        
        Args:
            messages: Messages to send to the model
            model: Model to use
            temperature: Temperature parameter for generation
            
        Returns:
            str: Analysis result
        """
        pass

class OpenAIAnalysisClient(AnalysisClient):
    """
    OpenAI implementation of AnalysisClient
    """
    def __init__(self, client=None):
        self.client = client or openai
    
    def analyze(self, messages, model="gpt-4-1106-preview", temperature=0):
        """
        Analyze text using OpenAI
        
        Args:
            messages: Messages to send to the model
            model: OpenAI model to use
            temperature: Temperature parameter for generation
            
        Returns:
            str: Analysis result
        """
        response = self.client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=messages
        )
        return response.choices[0].message.content

class TemplateSelector:
    """
    Selects the appropriate template for analysis
    """
    def __init__(self, prompt_templates, analysis_client):
        self.prompt_templates = prompt_templates
        self.analysis_client = analysis_client
    
    def select_template(self, text, **kwargs):
        """
        Auto-select the best template for the given text
        
        Args:
            text: Text to analyze
            **kwargs: Additional parameters
            
        Returns:
            dict: Selected template
        """
        template = self.prompt_templates.get_template("auto")
        try:
            messages = [
                {"role": "system", "content": template["system"]},
                {"role": "user", "content": template["template"].format(text=text)}
            ]
            
            analysis = self.analysis_client.analyze(messages)
            
            # Extract template name from the first line
            first_line = analysis.split('\n')[0].strip()
            recommended_template = first_line.split(':')[-1].strip().lower()
            recommended_template = recommended_template.replace('**', '').replace('*', '')
            
            # Log the template selection reasoning
            logger.info(f"Auto-selected template: {recommended_template}")
            logger.info(f"Selection reasoning: {analysis}")
            
            # Get and use the recommended template
            return self.prompt_templates.get_template(recommended_template, **kwargs)
        except Exception as e:
            logger.warning(f"Error in auto-template selection: {e}. Falling back to 'summary' template.")
            return self.prompt_templates.get_template("summary", **kwargs)

class MeetingAnalyzer:
    """
    Analyzes meeting transcriptions
    """
    def __init__(self, transcription: str, analysis_client=None, prompt_templates=None):
        """
        Initialize the analyzer
        
        Args:
            transcription: Text to analyze
            analysis_client: Client for analysis API
            prompt_templates: Templates for prompts
        """
        self.text_preprocessor = TextPreprocessor()
        self.transcription = self.text_preprocessor.prepare_text(transcription)
        self.analysis_client = analysis_client or OpenAIAnalysisClient()
        self.prompt_templates = prompt_templates or PromptTemplates()
        self.template_selector = TemplateSelector(self.prompt_templates, self.analysis_client)
    
    def analyze(self, template_name: str = "auto", **kwargs) -> str:
        """
        Analyze text using a specific template or auto-select the best template
        
        Args:
            template_name: Name of the template to use
            **kwargs: Additional parameters
            
        Returns:
            str: Analysis result
            
        Raises:
            AnalysisError: If analysis fails
        """
        try:
            if template_name == "auto":
                template = self.template_selector.select_template(self.transcription, **kwargs)
            else:
                template = self.prompt_templates.get_template(template_name, **kwargs)
                
            messages = [
                {"role": "system", "content": template["system"]},
                {"role": "user", "content": template["template"].format(text=self.transcription)}
            ]

            return self.analysis_client.analyze(messages)
        except openai.AuthenticationError as e:
            logger.error(f"Error en el análisis con template '{template_name}': {e}")
            raise AnalysisError(f"Authentication failed: {e}") from e
        except openai.APIError as e:
            logger.error(f"OpenAI API Error: {e}")
            raise AnalysisError(f"API Error: {e}") from e
        except Exception as e:
            logger.error(f"Unexpected error during analysis: {e}")
            raise AnalysisError(f"Unexpected error during analysis: {e}") from e

    def summarize(self, **kwargs):
        """
        Summarize the transcription
        
        Args:
            **kwargs: Additional parameters
            
        Returns:
            str: Summary
        """
        return self.analyze("summary", **kwargs)

    def extract_key_points(self, **kwargs):
        """
        Extract key points from the transcription
        
        Args:
            **kwargs: Additional parameters
            
        Returns:
            str: Key points
        """
        return self.analyze("key_points", **kwargs)

    def extract_action_items(self, **kwargs):
        """
        Extract action items from the transcription
        
        Args:
            **kwargs: Additional parameters
            
        Returns:
            str: Action items
        """
        return self.analyze("action_items", **kwargs)

    def analyze_sentiment(self, **kwargs):
        """
        Analyze sentiment in the transcription
        
        Args:
            **kwargs: Additional parameters
            
        Returns:
            str: Sentiment analysis
        """
        return self.analyze("sentiment", **kwargs)


class DocumentManager:
    """
    Manages document creation and saving
    """
    @staticmethod
    def create_document(content):
        """
        Create a document from content
        
        Args:
            content: Dictionary of content sections
            
        Returns:
            Document: Created document
        """
        doc = Document()
        for key, value in content.items():
            heading = key.replace('_', ' ').title()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        return doc
    
    @staticmethod
    def save_to_docx(content, filename):
        """
        Save content to a DOCX file
        
        Args:
            content: Dictionary of content sections
            filename: Path where to save the document
            
        Returns:
            str: Path to the saved document
        """
        # Ensure directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filename)), exist_ok=True)
        
        # Create and save document
        doc = DocumentManager.create_document(content)
        doc.save(filename)
        return filename


class DownloaderInterface:
    """
    Interface for video downloaders
    """
    def download(self, url, output_path=None):
        """
        Download a video from a URL
        
        Args:
            url: URL to download from
            output_path: Path where to save the video
            
        Returns:
            str: Path to the downloaded video
        """
        pass

class GoogleDriveDownloader(DownloaderInterface):
    """
    Downloads videos from Google Drive
    """
    def __init__(self, http_client=None):
        """
        Initialize the downloader
        
        Args:
            http_client: HTTP client for making requests
        """
        self.http_client = http_client or requests
    
    def download(self, drive_url, output_path=None):
        """
        Download a video from Google Drive
        
        Args:
            drive_url: Google Drive URL
            output_path: Path where to save the video
            
        Returns:
            str: Path to the downloaded video
            
        Raises:
            DownloadError: If download fails
        """
        try:
            file_id = drive_url.split('/')[-2]
            direct_download_url = f'https://drive.google.com/uc?id={file_id}&export=download'
            
            if not output_path:
                output_path = 'downloaded_video.mp4'
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            with self.http_client.get(direct_download_url, stream=True) as r:
                r.raise_for_status()
                with open(output_path, 'wb') as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
            
            return output_path
        except Exception as e:
            raise DownloadError(f"Failed to download from Google Drive: {e}") from e

class VideoDownloader:
    """
    Factory for video downloaders
    """
    @staticmethod
    def get_downloader(url):
        """
        Get the appropriate downloader for a URL
        
        Args:
            url: URL to download from
            
        Returns:
            DownloaderInterface: Downloader for the URL
        """
        if 'drive.google.com' in url:
            return GoogleDriveDownloader()
        # Add more downloaders as needed
        raise ValueError(f"No downloader available for URL: {url}")
    
    @staticmethod
    def download_from_google_drive(drive_url, output_path=None):
        """
        Download a video from Google Drive
        
        Args:
            drive_url: Google Drive URL
            output_path: Path where to save the video
            
        Returns:
            str: Path to the downloaded video
        """
        downloader = GoogleDriveDownloader()
        return downloader.download(drive_url, output_path)




def login_with_google():
    webbrowser.open('https://accounts.google.com/o/oauth2/auth?client_id=YOUR_CLIENT_ID&redirect_uri=YOUR_REDIRECT_URI&scope=https://www.googleapis.com/auth/drive.readonly&response_type=code', new=2)


def main(api_key, file_path, enable_diarization=False):
    try:
        logger.info("Beginning video transcription process...")

        # login_with_google()
        audio_file = AudioExtractor.extract_audio(file_path)
        logger.info("Audio extracted, starting transcription...")

        transcription_text = AudioTranscriptionService().transcribe(audio_file, diarization=enable_diarization)
        logger.info("Transcription completed.")

        logger.info("Analyzing transcription...")
        analyzer = MeetingAnalyzer(transcription_text)
        meeting_info = {
            'abstract_summary': analyzer.summarize(),
            'key_points': analyzer.extract_key_points(),
            'action_items': analyzer.extract_action_items(),
            'sentiment': analyzer.analyze_sentiment()
        }
        logger.info("Analysis completed.")

        logger.info("Saving analyzed information to document...")
        DocumentManager.save_to_docx(meeting_info, 'meeting_minutes.docx')
        logger.info("Document saved: meeting_minutes.docx")

        logger.info("\nMeeting Information:")
        for section, content in meeting_info.items():
            logger.info(f"{section.title().replace('_', ' ')}:\n{content}\n")

    except MeetingMinutesError as e:
        logger.error(f"MeetingMinutesError occurred: {e}")
        sys.exit(1)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        sys.exit(1)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description="Transcribe video to meeting minutes with optional speaker diarization"
    )
    parser.add_argument("api_key", type=str, help="API key for OpenAI")
    parser.add_argument("file_path", type=str, help="Path to the input video file")
    parser.add_argument("--diarization", action="store_true", help="Enable speaker diarization")
    args = parser.parse_args()

    main(args.api_key, args.file_path, args.diarization)
